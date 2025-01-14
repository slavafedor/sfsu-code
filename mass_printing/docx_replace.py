import argparse
from pathlib import Path
import re
import sys
import os
from typing import List
from docx import Document
import pandas as pd

def get_options(argv: List[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser()

    parser.add_argument(
        "-t", "--template-file",
        dest="templateFile",
        action="store", 
        required=True, 
        help="Set the input data file. This is a required argument." # If not supplied, the current folder will be used as a starting point."
    )

    parser.add_argument(
        "-d", "--data-file",
        dest="dataFile",
        action="store", 
        required=True, 
        help="Set the input data file. This is a required argument." # If not supplied, the current folder will be used as a starting point."
    )

    parser.add_argument(
        "-b", "--begin-index",
        dest="beginIndex",
        action="store", 
        type= int,
        default= None,
        help="Set the input data file starting data row index (starts with 0)." 
    )

    parser.add_argument(
        "-e", "--end-index",
        dest="endIndex",
        action="store", 
        type= int,
        default= None,
        help="Set the input data file ending data row index (starts with 0)." 
    )


    parser.add_argument(
        "-o", "--output-folder",
        dest="outputFolder",
        action="store",
        required=False, 
        default=Path.joinpath(Path.cwd(), "out"),
        help="Set the output folder. If not supplied, the <current folder>/out will be used as an output folder."
    )

    parser.add_argument(
        "-f", "--first-only",
        dest="firstOnly",
        action="store_true",
        required=False, 
        help="If set - the only first match will be replaced with value, and only one file will be created."
    )

    parser.add_argument(
        "-n", "--file-name-template",
        dest="fileNameTemplate",
        action="store",
        required=False, 
        default="{index}_{row[0]}_{row[1]}_{suffix}.docx",
        help="Set the output filename template. If not supplied, the default '{index}_{row[0]}_{row[1]}_{suffix}.docx' will be used."
    )

    options = parser.parse_args(argv)
    return options


def docx_replace_regex(doc_obj, regex, replace, first_only = False):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            # new_text = regex.sub(replace, p.text)
            # p.text = new_text
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):    
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
                    if first_only : return True

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                if docx_replace_regex(cell, regex, replace, first_only):
                    return

def main(argv: List[str] = sys.argv[1:]) -> None:
    options = get_options(argv)

    spreadsheet = options.dataFile
    template_name = options.templateFile
    suffix = os.path.splitext(os.path.basename(template_name))[0]
    firstOnly = options.firstOnly
    beginIndex = options.beginIndex
    endIndex = options.endIndex
    fileNameTemplate = options.fileNameTemplate

    if not os.path.exists(options.outputFolder):
        os.mkdir(options.outputFolder)

    print(f"spreadsheet = {spreadsheet}, template_name = {template_name}")

    excel_data = pd.read_excel(spreadsheet)
    excel_data = excel_data.iloc[beginIndex:endIndex]
    excel_data.fillna('', inplace=True)
    cols = excel_data.columns

    print(cols)
    for index, row in excel_data.iterrows():
        file_obj = Document(template_name)
        #if index != 29: continue
        for c in cols:
            if c is not None:
                regex1 = re.compile(f"{c}")
                val = "" if row[c] is None or row[c] == "nan" else str(row[c])
                print(f"c='{c}' : {val}")
                docx_replace_regex(file_obj, regex1, val, firstOnly)

        if not firstOnly:
            #file_name = os.path.join(options.outputFolder, f"{index}_{row[cols[0]]}_{row[cols[1]]}_{suffix}.docx")
            file_name = os.path.join(options.outputFolder, fileNameTemplate.format(**locals()))
        else: 
            file_name = os.path.join(options.outputFolder, f"out_{suffix}.docx")
            template_name = file_name

        print(file_name)
        file_obj.save(file_name)
        #exit()

if __name__ == "__main__":
    main()